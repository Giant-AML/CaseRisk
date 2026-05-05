#!/usr/bin/env python3
"""
Case-Specific Risk Assessment Generator
Giant Compliance Ltd | AML/CFT Compliance Tools

Fetches Companies House data and produces a pre-filled Word document
for the IP to complete the remaining manual checks.

Usage:
    python generate_case_ra.py
    python generate_case_ra.py --company 12345678
    python generate_case_ra.py --company 12345678 --api-key YOUR_KEY

Environment variable:
    CH_API_KEY  — your Companies House API key (avoids typing it each run)
"""

import os
import sys
import re
import argparse
import requests
from datetime import datetime, date
from dateutil import relativedelta
from io import BytesIO

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Colour palette ───────────────────────────────────────────────────────────
class C:
    NAVY       = "1F3864"
    MID_BLUE   = "2E75B6"
    LIGHT_BLUE = "BDD7EE"
    V_LIGHT    = "EEF4FB"
    GREEN_BG   = "C6EFCE";  GREEN_TX   = "375623"
    AMBER_BG   = "FFEB9C";  AMBER_TX   = "7D4C00"
    RED_BG     = "FFC7CE";  RED_TX     = "9C0006"
    WHITE      = "FFFFFF"
    GREY       = "F2F2F2"
    MED_GREY   = "BFBFBF"
    BLACK      = "000000"

# ─── Page geometry ────────────────────────────────────────────────────────────
# A4 with 1-inch margins; content width = 16.51 cm
MARGIN_CM  = 2.54   # 1 inch
PAGE_W_CM  = 21.0
CONTENT_CM = PAGE_W_CM - 2 * MARGIN_CM   # 15.92 cm

# ─── Companies House API ──────────────────────────────────────────────────────
CH_BASE = "https://api.companieshouse.gov.uk"

# ─── SIC code risk mapping ────────────────────────────────────────────────────
# Maps SIC prefixes to risk flags.  Keys are prefix strings; longest match wins.
SIC_RISK_MAP = {
    # Cash-intensive
    "47": "cash",       # Retail trade
    "56": "cash",       # Food & beverage
    "55": "cash",       # Accommodation
    "92": "cash",       # Gambling
    "93": "cash",       # Sports, leisure, amusement
    "96": "cash",       # Personal services (laundries, hairdressers, etc.)
    "45111": "cash",    # Sale of new cars
    "45112": "cash",    # Sale of used cars
    # High-value assets
    "641":  "hva",      # Banks / financial holding
    "642":  "hva",      # Financial holding companies
    "649":  "hva",      # Other credit
    "651":  "hva",      # Insurance
    "681":  "hva",      # Buying / selling real estate
    "682":  "hva",      # Letting / management of real estate
    "683":  "hva",      # Real estate agents
    "771":  "hva",      # Rental of vehicles
    "7711": "hva",
    "7712": "hva",
    # Cross-border
    "6420": "cross",    # Activities of holding companies
    "6430": "cross",    # Trusts, funds, similar financial entities
    # Regulated (FCA, SRA, etc.)
    "64":   "reg",      # Financial services
    "65":   "reg",      # Insurance
    "66":   "reg",      # Auxiliary financial services
    "6910": "reg",      # Legal activities (SRA)
    "6920": "reg",      # Accounting (ICAEW, ACCA)
    # Construction / property
    "41":   "prop",
    "42":   "prop",
    "43":   "prop",
    "68":   "prop",
    # Government contracting hint (public admin as SIC)
    "84":   "govt",
}

def sic_risk_flags(sic_codes: list[str]) -> dict[str, bool]:
    """Return a dict of risk flag -> True/False based on SIC codes."""
    flags = {"cash": False, "hva": False, "cross": False, "reg": False, "prop": False, "govt": False}
    for sic in sic_codes:
        sic_str = str(sic).strip()
        # Try longest prefix match first (5 chars down to 2)
        for length in [5, 4, 3, 2]:
            prefix = sic_str[:length]
            if prefix in SIC_RISK_MAP:
                flags[SIC_RISK_MAP[prefix]] = True
                break
    return flags

# ─── Companies House helpers ──────────────────────────────────────────────────

def ch_get(path: str, api_key: str) -> dict | None:
    """GET a CH API endpoint; returns parsed JSON or None on error."""
    url = CH_BASE + path
    try:
        r = requests.get(url, auth=(api_key, ""), timeout=10)
        if r.status_code == 200:
            return r.json()
        elif r.status_code == 404:
            return None
        else:
            print(f"  [CH] {r.status_code} for {path}")
            return None
    except requests.RequestException as e:
        print(f"  [CH] Request failed: {e}")
        return None


def fetch_company(company_number: str, api_key: str) -> dict:
    """Fetch and normalise company profile data."""
    data = ch_get(f"/company/{company_number}", api_key) or {}

    sic_codes = data.get("sic_codes", [])
    inc_date_raw = data.get("date_of_creation", "")
    inc_date = None
    trading_years = ""
    if inc_date_raw:
        try:
            inc_date = datetime.strptime(inc_date_raw, "%Y-%m-%d").date()
            delta = relativedelta.relativedelta(date.today(), inc_date)
            trading_years = f"{delta.years} years {delta.months} months"
        except ValueError:
            pass

    # Accounts info
    accts = data.get("accounts", {})
    next_due_raw   = accts.get("next_due", "")
    last_made_raw  = accts.get("last_accounts", {}).get("made_up_to", "")
    overdue        = accts.get("overdue", False)

    def fmt_date(raw):
        if not raw:
            return ""
        try:
            return datetime.strptime(raw, "%Y-%m-%d").strftime("%d/%m/%Y")
        except ValueError:
            return raw

    # Registered office
    roa = data.get("registered_office_address", {})
    address_parts = [
        roa.get("address_line_1", ""),
        roa.get("address_line_2", ""),
        roa.get("locality", ""),
        roa.get("postal_code", ""),
        roa.get("country", ""),
    ]
    address = ", ".join(p for p in address_parts if p)

    return {
        "name":           data.get("company_name", ""),
        "number":         data.get("company_number", company_number),
        "status":         data.get("company_status", ""),
        "type":           data.get("type", ""),
        "inc_date":       fmt_date(inc_date_raw),
        "trading_years":  trading_years,
        "sic_codes":      sic_codes,
        "address":        address,
        "accounts_overdue": overdue,
        "last_accounts":  fmt_date(last_made_raw),
        "next_accounts_due": fmt_date(next_due_raw),
    }


def fetch_officers(company_number: str, api_key: str) -> dict:
    """Return current and resigned directors/officers."""
    data = ch_get(f"/company/{company_number}/officers?items_per_page=100", api_key) or {}
    items = data.get("items", [])

    current, resigned = [], []
    for o in items:
        role = o.get("officer_role", "")
        name = o.get("name", "")
        appt = o.get("appointed_on", "")
        res  = o.get("resigned_on", "")
        nat  = o.get("nationality", "")
        dob_raw = o.get("date_of_birth", {})
        dob = ""
        if dob_raw:
            yr = dob_raw.get("year", "")
            mo = dob_raw.get("month", "")
            dob = f"{'0'+str(mo) if mo < 10 else mo}/{yr}" if yr and mo else str(yr)
        entry = {"name": name, "role": role, "appointed": appt, "resigned": res, "nationality": nat, "dob": dob}
        if res:
            resigned.append(entry)
        else:
            current.append(entry)

    return {"current": current, "resigned": resigned}


def fetch_pscs(company_number: str, api_key: str) -> list:
    """Return list of PSC entries."""
    data = ch_get(f"/company/{company_number}/persons-with-significant-control?items_per_page=100", api_key) or {}
    items = data.get("items", [])
    result = []
    for p in items:
        result.append({
            "name":        p.get("name", ""),
            "nationality": p.get("nationality", ""),
            "country":     p.get("country_of_residence", ""),
            "notified":    p.get("notified_on", ""),
            "controls":    ", ".join(p.get("natures_of_control", [])),
        })
    return result


def fetch_filing_count(company_number: str, api_key: str) -> dict:
    """Count total filings and accounts filings; flag late accounts."""
    data = ch_get(f"/company/{company_number}/filing-history?items_per_page=100", api_key) or {}
    items = data.get("items", [])
    total = data.get("total_count", len(items))
    accounts = [i for i in items if i.get("category") == "accounts"]
    return {
        "total":    total,
        "accounts": len(accounts),
        "latest_accounts_date": accounts[0].get("date", "") if accounts else "",
    }

def fetch_officer_appointments(name: str, api_key: str, subject_company: str = "") -> dict:
    """Search CH for an officer by name; return all their company appointments.

    Returns a dict:
      name, total, active, resigned, appointments (list of dicts), match_note, error
    """
    import urllib.parse
    result = {
        "name": name, "total": 0, "active": 0, "resigned": 0,
        "appointments": [], "match_note": "", "error": None,
    }
    try:
        q = urllib.parse.quote(name)
        data = ch_get(f"/search/officers?q={q}&items_per_page=20", api_key) or {}
        items = data.get("items", [])
        if not items:
            result["error"] = "No officer record found on CH"
            return result

        # Prefer exact name match; fall back to first result
        best = None
        for item in items:
            if item.get("title", "").upper().strip() == name.upper().strip():
                best = item
                break
        if not best:
            best = items[0]
            best_name = best.get("title", "")
            if best_name.upper() != name.upper():
                result["match_note"] = f"Closest CH match: {best_name}"

        # The links.self from officer search = /officers/{officer_id}
        self_link = best.get("links", {}).get("self", "")
        if not self_link:
            result["error"] = "No officer link available"
            return result

        appt_data = ch_get(f"{self_link}/appointments?items_per_page=50", api_key) or {}
        result["total"] = appt_data.get("total_results", 0)

        for item in appt_data.get("items", []):
            resigned = item.get("resigned_on", "")
            status   = item.get("appointed_to", {}).get("company_status", "")
            appt = {
                "company_name":   item.get("appointed_to", {}).get("company_name", ""),
                "company_number": item.get("appointed_to", {}).get("company_number", ""),
                "company_status": status or ("active" if not resigned else "unknown"),
                "role":           item.get("officer_role", "").replace("-", " ").title(),
                "appointed":      item.get("appointed_on", ""),
                "resigned":       resigned,
                "active":         not bool(resigned),
            }
            result["appointments"].append(appt)
            if not resigned:
                result["active"] += 1
            else:
                result["resigned"] += 1

        return result

    except Exception as e:
        result["error"] = str(e)
        return result

# ─── python-docx helpers ──────────────────────────────────────────────────────

def hex_to_rgb(hex_str: str):
    h = hex_str.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


def shade_cell(cell, hex_color: str):
    """Set background fill on a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove any existing shd
    for existing in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.upper())
    tcPr.append(shd)


def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    """Set internal cell padding in twips. Uses start/end (OOXML spec)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:tcMar")):
        tcPr.remove(existing)
    tcMar = OxmlElement("w:tcMar")
    # OOXML requires start/end (not left/right) for margins
    for side, val in (("top", top), ("start", left), ("bottom", bottom), ("end", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def set_cell_valign(cell, align="center"):
    """Set vertical alignment: 'top', 'center', 'bottom'."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:vAlign")):
        tcPr.remove(existing)
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), align)
    tcPr.append(vAlign)


def set_col_width(cell, width_cm: float):
    """Set a cell's preferred width. tcW must appear early in tcPr (before shd/vAlign)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:tcW")):
        tcPr.remove(existing)
    tcW = OxmlElement("w:tcW")
    # Convert cm to twips: 1 cm = 567 twips
    twips = int(width_cm * 567)
    tcW.set(qn("w:w"),    str(twips))
    tcW.set(qn("w:type"), "dxa")
    # Insert at position 0 so it precedes shd, tcMar, vAlign (schema order)
    tcPr.insert(0, tcW)


def write_cell(cell, text: str, bold=False, italic=False, size_pt=10,
               color_hex=None, align=WD_ALIGN_PARAGRAPH.LEFT,
               bg_hex=None, valign="center"):
    """Clear a cell and write formatted text into it."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size_pt)
    run.bold   = bold
    run.italic = italic
    if color_hex:
        r, g, b = hex_to_rgb(color_hex)
        run.font.color.rgb = RGBColor(r, g, b)
    if bg_hex:
        shade_cell(cell, bg_hex)
    set_cell_margins(cell)
    set_cell_valign(cell, valign)


def hdr_row(table, texts_widths: list[tuple], bg=C.NAVY, text_color=C.WHITE, size_pt=10):
    """Add a header row to a table. texts_widths: [(text, width_cm), ...]"""
    row = table.add_row()
    for i, (text, width) in enumerate(texts_widths):
        c = row.cells[i]
        write_cell(c, text, bold=True, size_pt=size_pt, color_hex=text_color, bg_hex=bg)
        set_col_width(c, width)
    return row


def sub_hdr_row(table, texts_widths: list[tuple]):
    """Light-blue sub-header row."""
    return hdr_row(table, texts_widths, bg=C.LIGHT_BLUE, text_color=C.BLACK)


def blank_row(table, widths_cm: list[float]):
    """Add a blank input row with correct column widths."""
    row = table.add_row()
    for i, w in enumerate(widths_cm):
        c = row.cells[i]
        write_cell(c, "", valign="top")
        set_col_width(c, w)
    return row


def section_banner(doc, num: int | str, title: str):
    """Full-width navy section header table."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    c = table.cell(0, 0)
    write_cell(c, f"SECTION {num}   |   {title.upper()}",
               bold=True, size_pt=12, color_hex=C.WHITE, bg_hex=C.NAVY)
    set_col_width(c, CONTENT_CM)
    doc.add_paragraph()  # spacer


def label_input_table(doc, rows: list[tuple], label_cm=5.5):
    """Two-column label / input table. rows: [(label, placeholder_text), ...]"""
    input_cm = CONTENT_CM - label_cm
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for label, hint in rows:
        row = table.add_row()
        lc = row.cells[0]
        ic = row.cells[1]
        write_cell(lc, label, bold=True, bg_hex=C.V_LIGHT)
        set_col_width(lc, label_cm)
        write_cell(ic, hint, italic=bool(hint), color_hex=C.MED_GREY if hint else None, valign="top")
        set_col_width(ic, input_cm)
    return table


def checklist_table(doc, items: list[str], label_cm=9.5):
    """Yes / No / N/A / Notes checklist table."""
    yn_cm  = 1.2
    na_cm  = 1.2
    not_cm = CONTENT_CM - label_cm - 2 * yn_cm - na_cm
    widths = [label_cm, yn_cm, yn_cm, na_cm, not_cm]

    table = doc.add_table(rows=0, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr_row(table,
            [("Risk / Check Item", label_cm), ("Yes", yn_cm), ("No", yn_cm),
             ("N/A", na_cm), ("Notes / Evidence", not_cm)],
            bg=C.LIGHT_BLUE, text_color=C.BLACK)
    for item in items:
        row = table.add_row()
        write_cell(row.cells[0], item, size_pt=9); set_col_width(row.cells[0], label_cm)
        for j in range(1, 5):
            write_cell(row.cells[j], "", valign="top")
            set_col_width(row.cells[j], widths[j])
    return table


def rag_selector_row(doc, label: str):
    """Single RAG selector row (Low / Medium / High + notes)."""
    lbl_cm = 5.5; rag_cm = 2.0; not_cm = CONTENT_CM - lbl_cm - 3 * rag_cm
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    r = table.rows[0]
    write_cell(r.cells[0], label, bold=True, bg_hex=C.V_LIGHT); set_col_width(r.cells[0], lbl_cm)
    write_cell(r.cells[1], "LOW",    bold=True, color_hex=C.GREEN_TX, bg_hex=C.GREEN_BG,
               align=WD_ALIGN_PARAGRAPH.CENTER); set_col_width(r.cells[1], rag_cm)
    write_cell(r.cells[2], "MEDIUM", bold=True, color_hex=C.AMBER_TX, bg_hex=C.AMBER_BG,
               align=WD_ALIGN_PARAGRAPH.CENTER); set_col_width(r.cells[2], rag_cm)
    write_cell(r.cells[3], "HIGH",   bold=True, color_hex=C.RED_TX,   bg_hex=C.RED_BG,
               align=WD_ALIGN_PARAGRAPH.CENTER); set_col_width(r.cells[3], rag_cm)
    write_cell(r.cells[4], "", valign="top"); set_col_width(r.cells[4], not_cm)
    return table


def narrative_box(doc, header: str, height_rows=3):
    """Full-width box with a sub-header and blank content rows."""
    table = doc.add_table(rows=0, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_row(table, [(header, CONTENT_CM)], bg=C.LIGHT_BLUE, text_color=C.BLACK)
    for _ in range(height_rows):
        r = table.add_row()
        write_cell(r.cells[0], "", valign="top")
        set_col_width(r.cells[0], CONTENT_CM)
    return table


def guidance(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run("Guidance: ")
    run.bold = True; run.italic = True; run.font.name = "Arial"; run.font.size = Pt(8.5)
    r, g, b = hex_to_rgb(C.MID_BLUE)
    run.font.color.rgb = RGBColor(r, g, b)
    run2 = p.add_run(text)
    run2.italic = True; run2.font.name = "Arial"; run2.font.size = Pt(8.5)
    run2.font.color.rgb = RGBColor(r, g, b)


def auto_tag(doc, text: str):
    """Small grey note to indicate the field was auto-populated."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(f"  Auto-populated from Companies House: {text}")
    r.font.name = "Arial"; r.font.size = Pt(7.5); r.italic = True
    r.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)


# ─── Pre-filled table helpers ─────────────────────────────────────────────────

def label_input_with_data(doc, rows_data: list[tuple], label_cm=5.5):
    """Like label_input_table but each row is (label, value, auto_filled_bool)."""
    input_cm = CONTENT_CM - label_cm
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for label, value, auto in rows_data:
        row = table.add_row()
        lc = row.cells[0]; ic = row.cells[1]
        write_cell(lc, label, bold=True, bg_hex=C.V_LIGHT)
        set_col_width(lc, label_cm)
        if auto and value:
            write_cell(ic, value, valign="top")
            # Add small auto-populated indicator in the same cell
            p2 = ic.add_paragraph()
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after  = Pt(0)
            rn = p2.add_run("[CH]")
            rn.font.name = "Arial"; rn.font.size = Pt(7); rn.italic = True
            rn.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
        else:
            write_cell(ic, value or "", italic=not value,
                       color_hex=C.MED_GREY if not value else None, valign="top")
        set_col_width(ic, input_cm)
    return table


def pre_filled_grid(doc, col_defs: list[dict], data_rows: list[list], extra_blank=2):
    """
    Generic grid with a header row + data rows.
    col_defs: [{"label": str, "width": float_cm}]
    data_rows: list of lists of cell values (str); auto-tag shown for non-empty cells.
    """
    widths = [cd["width"] for cd in col_defs]
    table = doc.add_table(rows=0, cols=len(col_defs))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_row(table, [(cd["label"], cd["width"]) for cd in col_defs],
            bg=C.LIGHT_BLUE, text_color=C.BLACK)
    for row_data in data_rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            c = row.cells[i]
            write_cell(c, str(val) if val else "", valign="top")
            if val:
                # Add [CH] tag
                p2 = c.add_paragraph()
                rn = p2.add_run("[CH]")
                rn.font.name = "Arial"; rn.font.size = Pt(7); rn.italic = True
                rn.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
            set_col_width(c, widths[i])
    # Extra blank rows
    for _ in range(extra_blank):
        blank_row(table, widths)
    return table


# ─── SIC description lookup ───────────────────────────────────────────────────

# Abbreviated lookup for common insolvency-sector codes
SIC_DESCRIPTIONS = {
    "41100": "Development of building projects",
    "41201": "Construction of commercial buildings",
    "41202": "Construction of domestic buildings",
    "43210": "Electrical installation",
    "43220": "Plumbing, heat & A/C installation",
    "43290": "Other construction installation",
    "43310": "Plastering",
    "43320": "Joinery installation",
    "43390": "Other building completion & finishing",
    "45111": "Sale of new cars & light motor vehicles",
    "45112": "Sale of used cars & light motor vehicles",
    "45190": "Sale of other motor vehicles",
    "45200": "Maintenance & repair of motor vehicles",
    "47110": "Retail in non-specialised stores with food",
    "47190": "Other retail in non-specialised stores",
    "47510": "Retail of textiles",
    "47710": "Retail of clothing",
    "47730": "Retail of pharmaceutical goods",
    "47789": "Other retail of new goods",
    "47910": "Retail via mail order / internet",
    "55100": "Hotels & similar accommodation",
    "55201": "Holiday centres & villages",
    "56101": "Licensed restaurants",
    "56102": "Unlicensed restaurants & cafes",
    "56210": "Event catering activities",
    "56302": "Public houses & bars",
    "58190": "Other publishing activities",
    "62012": "Business & domestic software development",
    "62090": "Other IT & computer services",
    "63990": "Other information service activities",
    "64110": "Central banking",
    "64191": "Banks",
    "64209": "Activities of other holding companies",
    "64922": "Activities of mortgage finance companies",
    "64999": "Other financial service activities",
    "65110": "Life insurance",
    "65120": "Non-life insurance",
    "66110": "Administration of financial markets",
    "66290": "Other activities auxiliary to insurance",
    "68100": "Buying & selling own real estate",
    "68201": "Renting & operating of housing association real estate",
    "68209": "Other letting & operating of own/leased real estate",
    "68310": "Real estate agencies",
    "68320": "Management of real estate on a fee or contract basis",
    "69102": "Solicitors",
    "69109": "Activities of patent & copyright agents; other legal activities",
    "69201": "Accounting & auditing activities",
    "69202": "Bookkeeping activities",
    "69203": "Tax consultancy",
    "70100": "Activities of head offices",
    "70229": "Management consultancy",
    "71111": "Architectural activities",
    "71112": "Urban planning & landscape activities",
    "71121": "Engineering design activities",
    "74909": "Other professional, scientific & technical activities",
    "77110": "Rental & leasing of cars & light motor vehicles",
    "77390": "Rental & leasing of other machinery, equipment & tangible goods",
    "82110": "Combined office administrative service activities",
    "82190": "Photocopying, document preparation & other specialised office activities",
    "82990": "Other business support service activities",
    "85590": "Other education",
    "86210": "General medical practice activities",
    "86900": "Other human health activities",
    "90010": "Performing arts",
    "92000": "Gambling & betting activities",
    "93110": "Operation of sports facilities",
    "93199": "Other sports activities",
    "96010": "Washing & (dry-)cleaning of textile & fur products",
    "96020": "Hairdressing & other beauty treatment",
    "96090": "Other personal service activities",
}

def sic_desc(code: str) -> str:
    code = str(code).strip().zfill(5)
    return SIC_DESCRIPTIONS.get(code, f"SIC {code}")


# ─── Document builder ─────────────────────────────────────────────────────────

def build_document(company_data: dict, officers: dict, pscs: list,
                   filing: dict, case_inputs: dict,
                   adverse_findings: list = None,
                   directorships: dict = None) -> Document:

    if adverse_findings is None:
        adverse_findings = []
    if directorships is None:
        directorships = {}

    # Build directorships summary string for the ch_rows table
    if directorships:
        parts = []
        for person, appts in directorships.items():
            active  = sum(1 for a in appts if a.get("active"))
            total   = len(appts)
            dissolved = sum(1 for a in appts if "dissolv" in a.get("company_status","").lower())
            insolv    = sum(1 for a in appts if any(
                w in a.get("company_status","").lower()
                for w in ("liquidat","administrat","receivership","voluntary")))
            note = f"{person}: {active} current, {total} total"
            if dissolved:
                note += f", {dissolved} dissolved"
            if insolv:
                note += f", {insolv} insolvent/in-procedure"
            parts.append(note)
        _directorships_summary = " | ".join(parts)
    else:
        _directorships_summary = "Not checked"

    doc = Document()

    # Page setup: A4, 1-inch margins
    section = doc.sections[0]
    section.page_width  = Cm(PAGE_W_CM)
    section.page_height = Cm(29.7)
    section.left_margin = section.right_margin = Cm(MARGIN_CM)
    section.top_margin  = section.bottom_margin = Cm(MARGIN_CM)

    # Header
    hdr_para = section.header.paragraphs[0]
    hdr_para.clear()
    hdr_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rh1 = hdr_para.add_run("[FIRM NAME]")
    rh1.bold = True; rh1.font.name = "Arial"; rh1.font.size = Pt(8)
    rh1.font.color.rgb = RGBColor(0xBF, 0xBF, 0xBF)
    rh2 = hdr_para.add_run("   |   CONFIDENTIAL  —  AML/CFT COMPLIANCE")
    rh2.font.name = "Arial"; rh2.font.size = Pt(8)
    rh2.font.color.rgb = RGBColor(0xBF, 0xBF, 0xBF)

    # Footer
    ftr = section.footer
    fp = ftr.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.clear()
    rftr = fp.add_run("Case-Specific Risk Assessment  |  Giant Compliance Ltd  |  CONFIDENTIAL")
    rftr.font.name = "Arial"; rftr.font.size = Pt(8)
    rftr.font.color.rgb = RGBColor(0xBF, 0xBF, 0xBF)

    # ── TITLE ──────────────────────────────────────────────────────────────────
    title_table = doc.add_table(rows=2, cols=1)
    title_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    write_cell(title_table.cell(0, 0),
               "CASE-SPECIFIC RISK ASSESSMENT",
               bold=True, size_pt=16, color_hex=C.WHITE, bg_hex=C.NAVY,
               align=WD_ALIGN_PARAGRAPH.CENTER)
    write_cell(title_table.cell(1, 0),
               "Anti-Money Laundering & Counter-Terrorist Financing  |  MLR 2017 (as amended)",
               bold=False, size_pt=10, color_hex=C.WHITE, bg_hex=C.MID_BLUE,
               align=WD_ALIGN_PARAGRAPH.CENTER)
    set_col_width(title_table.cell(0, 0), CONTENT_CM)
    set_col_width(title_table.cell(1, 0), CONTENT_CM)
    doc.add_paragraph()

    # ── PART A — CASE DETAILS ─────────────────────────────────────────────────
    p = doc.add_heading("PART A — CASE DETAILS", level=1)
    p.runs[0].font.name = "Arial"; p.runs[0].font.size = Pt(12)
    p.runs[0].font.color.rgb = RGBColor(*hex_to_rgb(C.NAVY))

    guidance(doc, "Complete all fields before proceeding. [CH] = auto-populated from Companies House. "
                  "Verify auto-populated data and amend if required.")

    today_str = date.today().strftime("%d/%m/%Y")
    sic_list  = company_data.get("sic_codes", [])
    sic_display = "; ".join(f"{c} — {sic_desc(c)}" for c in sic_list) if sic_list else ""

    case_rows = [
        ("Case Reference",               case_inputs.get("case_ref", ""),        False),
        ("Case Name / Trading Name",     company_data.get("name", ""),           True),
        ("Company Registration Number",  company_data.get("number", ""),         True),
        ("Appointment Type",             case_inputs.get("appt_type", "CVL / Administration / MVL / Other"), False),
        ("Date of Appointment",          case_inputs.get("appt_date", ""),       False),
        ("IP / Officeholder Name",       case_inputs.get("ip_name", ""),         False),
        ("IP Licence Number",            case_inputs.get("ip_licence", ""),      False),
        ("SIC Code(s) & Industry",       sic_display,                             True),
        ("Date of Assessment",           today_str,                              False),
        ("Assessed By",                  case_inputs.get("assessed_by", ""),     False),
        ("Next Review Date",             "DD/MM/YYYY",                           False),
    ]
    label_input_with_data(doc, case_rows, label_cm=5.5)
    doc.add_paragraph()

    # ── SECTION 1 — COMPANIES HOUSE SUMMARY ───────────────────────────────────
    section_banner(doc, 1, "Companies House Background Summary")
    guidance(doc, "Auto-populated fields are drawn from the CH API and should be verified. "
                  "Add observations in the narrative box below.")

    # Last accounts info
    last_accts = company_data.get("last_accounts", "")
    next_due   = company_data.get("next_accounts_due", "")
    overdue    = company_data.get("accounts_overdue", False)
    overdue_str = "YES — accounts overdue at Companies House" if overdue else "No"

    # Current directors
    cur_dirs = "; ".join(
        f"{o['name']} (appt {o['appointed']}, {o['role']})"
        for o in officers.get("current", [])[:10]
    ) or "See CH filing history"

    # Former directors
    res_dirs = "; ".join(
        f"{o['name']} (res {o['resigned']})"
        for o in officers.get("resigned", [])[:10]
    ) or "None noted / see CH"

    # PSCs
    psc_str = "; ".join(
        f"{p['name']} ({p['controls'][:60]})" for p in pscs[:5]
    ) or "None registered / see CH"

    ch_rows = [
        ("Date of Incorporation",          company_data.get("inc_date", ""),           True),
        ("Period of Trading (approx.)",    company_data.get("trading_years", ""),       True),
        ("SIC Code(s) & Industry",         sic_display,                                True),
        ("Registered Office Address",      company_data.get("address", ""),             True),
        ("Company Status",                 company_data.get("status", "").replace("-", " ").title(), True),
        ("Number of Accounts Filed (CH sample)", str(filing.get("accounts", "")),      True),
        ("Date of Latest Accounts",        last_accts,                                  True),
        ("Next Accounts Due",              next_due,                                    True),
        ("Overdue / Late Filings?",        overdue_str,                                 True),
        ("Registered Charges / Mortgages", "Check CH charges register",                False),
        ("Current Director(s)",            cur_dirs,                                    True),
        ("Former Director(s)",             res_dirs,                                    True),
        ("Persons of Significant Control", psc_str,                                     True),
        ("Share Capital & Ownership",      "See CH filing — confirm with latest accounts", False),
        ("Associated / Related Companies", _directorships_summary, bool(_directorships_summary and _directorships_summary != "Not checked")),
        ("CH Flags / Concerns",            "e.g. late filings, dormancy gaps, frequent officer changes", False),
    ]
    label_input_with_data(doc, ch_rows, label_cm=5.5)
    doc.add_paragraph()
    narrative_box(doc, "CH Summary — Narrative Observations (key concerns, flags, background)", height_rows=4)
    doc.add_paragraph()

    # ── Directorships detail table ─────────────────────────────────────────────
    if directorships:
        guidance(doc,
            "Cross-directorships sourced from Companies House. Current appointments highlighted. "
            "Multiple dissolved or insolvent companies are a risk indicator requiring narrative explanation.")
        dir_cols = [
            {"label": "Principal",      "width": 3.0},
            {"label": "Company",        "width": 4.5},
            {"label": "Co. No.",        "width": 1.4},
            {"label": "Status",         "width": 1.8},
            {"label": "Role",           "width": 2.0},
            {"label": "Appointed",      "width": 1.6},
            {"label": "Resigned",       "width": 1.62},
        ]
        # 3.0+4.5+1.4+1.8+2.0+1.6+1.62 = 15.92
        dir_widths = [cd["width"] for cd in dir_cols]
        dir_table = doc.add_table(rows=0, cols=len(dir_cols))
        dir_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr_row(dir_table, [(cd["label"], cd["width"]) for cd in dir_cols],
                bg=C.LIGHT_BLUE, text_color=C.BLACK)
        for person, appt_list in directorships.items():
            for appt in appt_list:
                is_active = appt.get("active", False)
                row = dir_table.add_row()
                vals = [
                    person,
                    appt.get("company_name", ""),
                    appt.get("company_number", ""),
                    appt.get("company_status", "").replace("-", " ").title(),
                    appt.get("role", ""),
                    appt.get("appointed", ""),
                    appt.get("resigned", "") or "Current",
                ]
                for i, (val, wd) in enumerate(zip(vals, dir_widths)):
                    c = row.cells[i]
                    bg = C.LIGHT_BLUE if is_active else None
                    write_cell(c, val, size_pt=9, valign="top", bg_hex=bg)
                    set_col_width(c, wd)
        doc.add_paragraph()

    # ── SECTION 2 — SECTOR RISK ───────────────────────────────────────────────
    section_banner(doc, 2, "Industry Sector Risk Assessment")
    guidance(doc, "Risk flags below are auto-suggested based on SIC codes. Review each flag and amend as required for this specific case.")

    flags = sic_risk_flags(sic_list)

    # Sector profile table
    geo_scope = "Local / National / International — detail any international connections"
    label_input_with_data(doc, [
        ("SIC Code(s) / Industry",         sic_display,  True),
        ("Nature of Business",             "Describe what the company did and how it operated", False),
        ("Geographic Scope",               geo_scope,    False),
        ("Typical Transaction Profile",    "Cash / card / BACS / trade credit / crypto / other", False),
    ], label_cm=5.5)
    doc.add_paragraph()

    # Checklist with auto-flags
    yn_cm = 1.2; na_cm = 1.2
    label_cm_cl = 9.5
    not_cm = CONTENT_CM - label_cm_cl - 2 * yn_cm - na_cm
    widths_cl = [label_cm_cl, yn_cm, yn_cm, na_cm, not_cm]

    checklist_items = [
        ("Cash-intensive business (retail, hospitality, leisure, car wash, etc.)", flags["cash"]),
        ("High-value physical assets (jewellery, art, vehicles, property)",        flags["hva"]),
        ("Cross-border or international payments / operations",                     flags["cross"]),
        ("Regulated sector (FCA, SRA, CQC, HMRC authorised, etc.)",               flags["reg"]),
        ("Government or public sector contracting",                                 flags["govt"]),
        ("Construction, property development or land transactions",                 flags["prop"]),
        ("Complex or multi-layered corporate structure",                            False),
        ("Digital assets / cryptocurrency involvement",                             False),
        ("Sector subject to known NCA / UKFIU typology warnings",                  False),
        ("Significant HMRC / Crown debt (indicates higher sector risk)",            False),
        ("Prior sector regulatory action / enforcement",                            False),
    ]

    cl_table = doc.add_table(rows=0, cols=5)
    cl_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_row(cl_table,
            [("Risk / Check Item", label_cm_cl), ("Yes", yn_cm), ("No", yn_cm),
             ("N/A", na_cm), ("Notes / Evidence", not_cm)],
            bg=C.LIGHT_BLUE, text_color=C.BLACK)

    for item_text, auto_yes in checklist_items:
        row = cl_table.add_row()
        write_cell(row.cells[0], item_text, size_pt=9)
        set_col_width(row.cells[0], label_cm_cl)
        for j, (ww, pre) in enumerate([(yn_cm, auto_yes), (yn_cm, False), (na_cm, False)]):
            val = "[CH]" if (j == 0 and auto_yes) else ""
            bg  = C.AMBER_BG if (j == 0 and auto_yes) else C.WHITE
            write_cell(row.cells[j+1], val, size_pt=8, italic=True,
                       color_hex=C.AMBER_TX if val else None, bg_hex=bg,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
            set_col_width(row.cells[j+1], ww)
        write_cell(row.cells[4], "", valign="top"); set_col_width(row.cells[4], not_cm)

    doc.add_paragraph()
    rag_selector_row(doc, "Sector Risk Rating (circle / highlight applicable)")
    doc.add_paragraph()

    # ── SECTION 3 — CDD ON PRINCIPALS ────────────────────────────────────────
    section_banner(doc, 3, "Customer Due Diligence (CDD) on Principals")
    guidance(doc, "Director and PSC names are pre-populated from Companies House. "
                  "Complete ID verification, address verification and date for each principal. "
                  "Attach copies of ID documents to the case file.")

    cdd_cols = [
        {"label": "Principal Name",        "width": 3.2},
        {"label": "Role",                  "width": 2.0},
        {"label": "DOB",                   "width": 1.5},
        {"label": "Nationality",           "width": 1.8},
        {"label": "ID Obtained (Y/N)",     "width": 1.5},
        {"label": "ID Type / Reference",   "width": 2.5},
        {"label": "Address Verified (Y/N)","width": 1.7},
        {"label": "Date Checked",          "width": 1.72},
    ]
    # Total: 3.2+2.0+1.5+1.8+1.5+2.5+1.7+1.72 = 15.92 = CONTENT_CM ✓

    cdd_widths = [cd["width"] for cd in cdd_cols]

    cdd_table = doc.add_table(rows=0, cols=len(cdd_cols))
    cdd_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_row(cdd_table, [(cd["label"], cd["width"]) for cd in cdd_cols],
            bg=C.LIGHT_BLUE, text_color=C.BLACK)

    # Pre-populate from directors + PSCs
    principals = []
    for o in officers.get("current", []):
        principals.append({"name": o["name"], "role": o["role"].replace("-", " ").title(),
                            "dob": o.get("dob", ""), "nat": o.get("nationality", "")})
    for p in pscs:
        principals.append({"name": p["name"], "role": "PSC",
                            "dob": "", "nat": p.get("nationality", "")})

    for pr in principals[:10]:
        row = cdd_table.add_row()
        vals = [pr["name"], pr["role"], pr["dob"], pr["nat"], "", "", "", ""]
        for i, (val, wd) in enumerate(zip(vals, cdd_widths)):
            c = row.cells[i]
            write_cell(c, val, valign="top")
            if val and i < 4:
                p2 = c.add_paragraph()
                rn = p2.add_run("[CH]")
                rn.font.name = "Arial"; rn.font.size = Pt(7); rn.italic = True
                rn.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
            set_col_width(c, wd)

    # 3 extra blank rows
    for _ in range(3):
        blank_row(cdd_table, cdd_widths)

    doc.add_paragraph()
    narrative_box(doc, "CDD Notes — UBO / Beneficial Ownership / Additional Observations", height_rows=3)
    doc.add_paragraph()

    # ── SECTION 4 — PEP & SANCTIONS ──────────────────────────────────────────
    section_banner(doc, 4, "PEP and Sanctions Screening")
    guidance(doc, "Screen all principals against PEP lists and consolidated sanctions lists "
                  "(OFSI, UN, EU, OFAC as appropriate). Record date, tool used and outcome. "
                  "Names are pre-populated from Companies House — add any additional principals.")

    label_input_with_data(doc, [
        ("Screening Tool / Database Used", "e.g. Creditsafe, Comply Advantage, World-Check, manual check", False),
        ("Date Screening Undertaken",      "", False),
    ], label_cm=5.5)
    doc.add_paragraph()

    pep_cols = [
        {"label": "Principal Name",       "width": 3.2},
        {"label": "PEP Result",           "width": 1.8},
        {"label": "Sanctions Hit",        "width": 1.8},
        {"label": "Adverse Hit",          "width": 1.8},
        {"label": "Cleared / Resolved?",  "width": 1.8},
        {"label": "Notes / Action Taken", "width": 5.52},
    ]
    # 3.2+1.8+1.8+1.8+1.8+5.52 = 15.92 ✓

    pep_widths = [cd["width"] for cd in pep_cols]
    pep_table = doc.add_table(rows=0, cols=len(pep_cols))
    pep_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_row(pep_table, [(cd["label"], cd["width"]) for cd in pep_cols],
            bg=C.LIGHT_BLUE, text_color=C.BLACK)

    all_principals = [(o["name"], o["role"].replace("-", " ").title()) for o in officers.get("current", [])]
    all_principals += [(p["name"], "PSC") for p in pscs]

    for name, role in all_principals[:12]:
        row = pep_table.add_row()
        for i, (val, wd) in enumerate(zip([name, "", "", "", "", ""], pep_widths)):
            c = row.cells[i]
            write_cell(c, val, valign="top")
            if val and i == 0:
                p2 = c.add_paragraph()
                rn = p2.add_run("[CH]")
                rn.font.name = "Arial"; rn.font.size = Pt(7); rn.italic = True
                rn.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
            set_col_width(c, wd)
    for _ in range(2):
        blank_row(pep_table, pep_widths)

    doc.add_paragraph()
    narrative_box(doc, "PEP / Sanctions Narrative (explain any hits, near-matches or enhanced checks)", height_rows=3)
    doc.add_paragraph()

    # -- SECTION 5 -- ADVERSE MEDIA & SANCTIONS SCREENING ---------------------
    principals_str = "; ".join(
        o["name"] for o in officers.get("current", [])[:8]
    ) + ("; ".join(p["name"] for p in pscs[:4]) or "")
    if not principals_str:
        principals_str = "See Sections 3 & 4"

    section_banner(doc, 5, "Adverse Media and Sanctions Screening")
    guidance(doc,
        "Record the URL or document reference for EVERY finding — this is the audit trail. "
        "If a concern is identified, complete the Evidence Log (Part D) with the link. "
        "Even a 'clear' result should be recorded with the date checked.")

    # Part A: High-Risk Indicator Flags
    hi_flag_label = 3.5
    notes_table = doc.add_table(rows=0, cols=1)
    notes_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_row(notes_table, [("Part A  |  High-Risk Indicator Flags", CONTENT_CM)],
            bg=C.MID_BLUE, text_color=C.WHITE)
    doc.add_paragraph()
    hi_items = [
        "Named on an official sanctions list: UK OFSI / US OFAC / UN Consolidated / EU",
        "Subject of a US Senate, UK parliamentary or other governmental inquiry or published report",
        "Named in major investigative journalism: ICIJ (Panama Papers, FinCEN Files, Pandora Papers), Bureau of Investigative Journalism",
        "Significant adverse coverage in mainstream financial press: Financial Times, Reuters, Bloomberg",
        "Referenced in Wikileaks, document leaks or whistleblower disclosures",
        "Wikipedia entry contains substantive adverse, controversy or criminal content",
        "Court proceedings, tribunal decisions, regulatory enforcement or criminal convictions documented",
        "Disqualified director or subject of insolvency-related restriction (Insolvency Service register)",
        "Known connection to sanctioned individuals, Politically Exposed Persons or criminal networks identified in media",
        "Adverse content in BBC, Guardian, national or regional press suggesting financial misconduct",
    ]
    checklist_table(doc, hi_items, label_cm=9.5)
    doc.add_paragraph()

    # Helper for source tables (same columns throughout)
    S_COLS = [
        {"label": "Source / Database",         "width": 3.2},
        {"label": "Persons / Entity Searched",  "width": 2.8},
        {"label": "Date",                       "width": 1.4},
        {"label": "Finding / Reference",        "width": 4.5},
        {"label": "URL / Document Link",        "width": 2.62},
        {"label": "Concern",                    "width": 1.4},
    ]
    # 3.2+2.8+1.4+4.5+2.62+1.4 = 15.92 CHECK
    s_widths = [cd["width"] for cd in S_COLS]

    def am_source_table(items_with_data):
        """Build a source-row table. items_with_data: list of (source_name, prefill_finding, prefill_url)."""
        tbl = doc.add_table(rows=0, cols=len(S_COLS))
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr_row(tbl, [(cd["label"], cd["width"]) for cd in S_COLS],
                bg=C.LIGHT_BLUE, text_color=C.BLACK)
        for src, finding, url in items_with_data:
            # Get any pre-filled data from adverse_findings
            af = {f["source"]: f for f in adverse_findings}
            match = af.get(src, {})
            row = tbl.add_row()
            vals = [
                src,
                match.get("persons", principals_str[:60]),
                match.get("date", ""),
                match.get("finding", finding or ""),
                match.get("url", url or ""),
                match.get("concern", ""),
            ]
            for i, (val, wd) in enumerate(zip(vals, s_widths)):
                c = row.cells[i]
                is_auto = bool(match) and i in (1, 2, 3, 4, 5)
                write_cell(c, val, size_pt=8 if i == 0 else 9, valign="top",
                           italic=(i == 0))
                if is_auto and val:
                    p2 = c.add_paragraph()
                    rn = p2.add_run("[APP]")
                    rn.font.name = "Arial"; rn.font.size = Pt(7); rn.italic = True
                    rn.font.color.rgb = RGBColor(0x37, 0x56, 0x23)
                set_col_width(c, wd)
        return tbl

    # Part B: Sanctions & Official Lists
    hdr_row(doc.add_table(rows=1, cols=1).rows[0].cells[0].__class__.__mro__  # get a banner without calling section_banner
            if False else None, [], bg=C.NAVY) if False else None
    ban_b = doc.add_table(rows=1, cols=1)
    ban_b.alignment = WD_TABLE_ALIGNMENT.LEFT
    write_cell(ban_b.cell(0,0), "Part B  |  Sanctions and Official Lists",
               bold=True, size_pt=11, color_hex=C.WHITE, bg_hex=C.NAVY)
    set_col_width(ban_b.cell(0,0), CONTENT_CM)
    doc.add_paragraph()

    am_source_table([
        ("OFSI  —  UK HM Treasury Consolidated Sanctions List",
         "", "https://www.gov.uk/government/publications/financial-sanctions-consolidated-list-of-targets"),
        ("OFAC  —  US Treasury Office of Foreign Assets Control",
         "", "https://sanctionssearch.ofac.treas.gov/"),
        ("UN Security Council Consolidated Sanctions List",
         "", "https://www.un.org/securitycouncil/content/un-sc-consolidated-list"),
        ("EU Consolidated Sanctions (Sanctions Map)",
         "", "https://www.sanctionsmap.eu/"),
        ("HMRC Published Customs Notices / Excise Approvals",
         "", "https://www.gov.uk/government/organisations/hm-revenue-customs"),
    ])
    doc.add_paragraph()

    # Part C: Official Registers & Legal Records
    ban_c = doc.add_table(rows=1, cols=1)
    ban_c.alignment = WD_TABLE_ALIGNMENT.LEFT
    write_cell(ban_c.cell(0,0), "Part C  |  Official Registers and Legal Records",
               bold=True, size_pt=11, color_hex=C.WHITE, bg_hex=C.NAVY)
    set_col_width(ban_c.cell(0,0), CONTENT_CM)
    doc.add_paragraph()

    am_source_table([
        ("Companies House  —  Register of Disqualified Directors",
         "", "https://find-and-update.company-information.service.gov.uk/register-of-disqualifications/A"),
        ("Insolvency Service  —  Bankruptcy / IVA / DRO Register",
         "", "https://www.insolvencydirect.bis.gov.uk/eiir/"),
        ("FCA  —  Financial Services Register (individuals & firms)",
         "", "https://register.fca.org.uk/"),
        ("Find Case Law / BAILII  —  Court Judgments & Tribunal Decisions",
         "", "https://caselaw.nationalarchives.gov.uk/"),
        ("The Gazette  —  Insolvency and Winding-Up Notices",
         "", "https://www.thegazette.co.uk/"),
    ])
    doc.add_paragraph()

    # Part D: Media & Investigative Journalism
    ban_d = doc.add_table(rows=1, cols=1)
    ban_d.alignment = WD_TABLE_ALIGNMENT.LEFT
    write_cell(ban_d.cell(0,0), "Part D  |  Media and Investigative Journalism",
               bold=True, size_pt=11, color_hex=C.WHITE, bg_hex=C.NAVY)
    set_col_width(ban_d.cell(0,0), CONTENT_CM)
    doc.add_paragraph()

    company_name = company_data.get("name", "")
    am_source_table([
        ("Google News  —  Company and all directors / PSCs",
         "", "https://news.google.com/search?q=" + company_name.replace(" ", "+")),
        ("Financial Times",
         "", "https://www.ft.com/search?q=" + company_name.replace(" ", "+")),
        ("Bureau of Investigative Journalism",
         "", "https://www.thebureauinvestigates.com/?s=" + company_name.replace(" ", "+")),
        ("ICIJ Offshore Leaks Database (Panama Papers / FinCEN Files / Pandora Papers)",
         "", "https://offshoreleaks.icij.org/search?q=" + company_name.replace(" ", "+")),
        ("Wikipedia  —  Directors and company (check controversy / biography sections)",
         "", "https://en.wikipedia.org/wiki/Special:Search?search=" + company_name.replace(" ", "+")),
        ("Wikileaks  —  Principals and company name",
         "", "https://search.wikileaks.org/?q=" + company_name.replace(" ", "+")),
        ("BBC / Reuters / Bloomberg / national press",
         "", ""),
        ("Local / trade press  —  sector and geographic area",
         "", ""),
    ])
    doc.add_paragraph()

    # Part E: Evidence Log (URLs for all concerns)
    ban_e = doc.add_table(rows=1, cols=1)
    ban_e.alignment = WD_TABLE_ALIGNMENT.LEFT
    write_cell(ban_e.cell(0,0),
               "Part E  |  Evidence Log  (record every URL / document reference for concerns found)",
               bold=True, size_pt=11, color_hex=C.WHITE, bg_hex=C.MID_BLUE)
    set_col_width(ban_e.cell(0,0), CONTENT_CM)
    doc.add_paragraph()

    guidance(doc, "Paste the full URL or document reference for each concern identified. "
                  "If the source is a physical document (Senate report, court bundle, etc.), "
                  "note the title, date and where it is stored. This log is the primary evidence record.")

    ev_cols = [
        {"label": "URL / Document Reference",      "width": 5.5},
        {"label": "Date Accessed",                 "width": 1.5},
        {"label": "Person / Entity Concerned",     "width": 2.8},
        {"label": "Source",                        "width": 2.0},
        {"label": "Finding / Concern Summary",     "width": 2.72},
        {"label": "Risk",                          "width": 1.4},
    ]
    # 5.5+1.5+2.8+2.0+2.72+1.4 = 15.92 CHECK
    ev_widths = [cd["width"] for cd in ev_cols]
    ev_table = doc.add_table(rows=0, cols=len(ev_cols))
    ev_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_row(ev_table, [(cd["label"], cd["width"]) for cd in ev_cols],
            bg=C.LIGHT_BLUE, text_color=C.BLACK)

    # Pre-populate evidence log from adverse_findings where concern = Y
    ev_pre = [f for f in adverse_findings if f.get("concern", "").upper() == "Y"]
    for f in ev_pre:
        row = ev_table.add_row()
        for i, (val, wd) in enumerate(zip([
            f.get("url",""), "", f.get("persons",""), f.get("source",""),
            f.get("finding",""), "HIGH"
        ], ev_widths)):
            c = row.cells[i]
            write_cell(c, val, size_pt=9, valign="top",
                       color_hex=(C.RED_TX if i==5 else None),
                       bg_hex=(C.RED_BG if i==5 else None))
            set_col_width(c, wd)

    for _ in range(max(6, 6 - len(ev_pre))):
        blank_row(ev_table, ev_widths)
    doc.add_paragraph()

    # Part F: Overall assessment
    rag_selector_row(doc, "Adverse Media Overall Risk Rating")
    doc.add_paragraph()
    narrative_box(doc,
        "Adverse Media Narrative  —  Summary of Findings, Risk Conclusion and Any Outstanding Actions",
        height_rows=4)
    doc.add_paragraph()


        # ── SECTION 6 — SOURCE OF WEALTH / FUNDS ─────────────────────────────────
    section_banner(doc, 6, "Source of Wealth and Source of Funds")
    guidance(doc, "Assess the origin of assets in the estate and financial flows into the business. "
                  "Obtain documentary evidence where possible. Record any concerns or anomalies.")

    label_input_with_data(doc, [
        ("Nature of Assets in the Estate",         "e.g. plant & machinery, stock, property, book debts, cash", False),
        ("Estimated Gross Asset Value",             "",   False),
        ("Source of Wealth — Director Explanation","How directors explain origin of business assets / funds", False),
        ("SOW Documentary Evidence Obtained?",     "Yes / No / Partial — specify documents", False),
        ("SOW Documents Obtained",                 "e.g. bank statements, valuation reports, title deeds", False),
        ("Source of Funds — Explanation",          "Origin of cash / financial flows into the estate", False),
        ("SOF Evidence Obtained?",                 "Yes / No / Partial", False),
        ("SOF Documents",                          "e.g. bank statements, HMRC records, loan agreements", False),
        ("Concerns / Anomalies Noted",             "", False),
    ], label_cm=5.5)
    doc.add_paragraph()
    rag_selector_row(doc, "Source of Wealth / Funds Risk Rating")
    doc.add_paragraph()

    # ── SECTION 7 — INSOLVENCY RISK FLAGS ────────────────────────────────────
    section_banner(doc, 7, "Initial Insolvency Risk Flags")
    guidance(doc, "Work through each flag at the initial stages of the appointment. "
                  "Update this section as the case progresses and new information emerges.")

    checklist_table(doc, [
        "Phoenix activity / pre-pack indicators (same business restarted by same principals)",
        "Overdrawn Director's Loan Account (DLA) requiring investigation",
        "Potential antecedent transactions — preferences (s.239 IA 1986)",
        "Potential antecedent transactions — transactions at undervalue (s.238 IA 1986)",
        "Transactions defrauding creditors (s.423 IA 1986) indicators",
        "Missing, incomplete or falsified books and records",
        "Significant HMRC / Crown debt (PAYE, VAT, CT arrears)",
        "Connected party transactions (sales to directors, family, related entities)",
        "Unusual asset movements or disposals shortly before insolvency",
        "International assets or connections (offshore accounts, overseas property)",
        "Potential undisclosed or hidden assets",
        "Prior insolvency history — director personally or other companies",
        "Litigation, regulatory action or enforcement (company or directors)",
        "Complex or opaque ownership / corporate structure",
        "Creditor fraud indicators (false invoicing, fictitious creditors)",
        "Personal guarantees provided — guarantor risk",
        "Employee / payroll irregularities",
        "Wrongful trading indicators (s.214 IA 1986)",
    ], label_cm=9.5)

    doc.add_paragraph()
    narrative_box(doc, "Potential Assets to Investigate", height_rows=1)

    # Asset investigation table
    asset_cols = [
        {"label": "Asset / Investigation Area", "width": 4.5},
        {"label": "Est. Value",                 "width": 2.0},
        {"label": "Initial Evidence / Source",  "width": 5.42},
        {"label": "Action Required",            "width": 4.0},
    ]
    # 4.5+2.0+5.42+4.0 = 15.92 ✓
    asset_widths = [cd["width"] for cd in asset_cols]
    asset_table = doc.add_table(rows=0, cols=4)
    asset_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_row(asset_table, [(cd["label"], cd["width"]) for cd in asset_cols],
            bg=C.LIGHT_BLUE, text_color=C.BLACK)
    for _ in range(6):
        blank_row(asset_table, asset_widths)

    doc.add_paragraph()

    # ── SECTION 8 — OVERALL RISK RATING ──────────────────────────────────────
    section_banner(doc, 8, "Overall Risk Rating and Decisions")
    guidance(doc, "Aggregate the risk dimensions below. Circle / highlight the applicable rating "
                  "for each dimension, then record the overall case rating and decisions.")

    # Risk summary table
    dim_cols = [3.5, 1.8, 1.8, 1.8, 7.02]
    # 3.5+1.8+1.8+1.8+7.02 = 15.92 ✓
    dim_table = doc.add_table(rows=0, cols=5)
    dim_table.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr_row(dim_table,
            [("Risk Dimension", dim_cols[0]),
             ("LOW",    dim_cols[1]),
             ("MEDIUM", dim_cols[2]),
             ("HIGH",   dim_cols[3]),
             ("Rationale / Notes", dim_cols[4])],
            bg=C.LIGHT_BLUE, text_color=C.BLACK)

    dimensions = [
        "Sector / Industry Risk",
        "Companies House / Background Risk",
        "CDD / KYC Adequacy",
        "PEP & Sanctions Risk",
        "Adverse Media Risk",
        "Source of Wealth / Funds Risk",
        "Insolvency Risk Flags",
        "Geographic Risk",
        "Complexity / Structure Risk",
    ]
    for dim in dimensions:
        row = dim_table.add_row()
        write_cell(row.cells[0], dim, size_pt=9); set_col_width(row.cells[0], dim_cols[0])
        for j, (wd, (bg, tx)) in enumerate(zip(dim_cols[1:4], [
            (C.GREEN_BG, C.GREEN_TX), (C.AMBER_BG, C.AMBER_TX), (C.RED_BG, C.RED_TX)
        ])):
            write_cell(row.cells[j+1], "", bg_hex=bg, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_col_width(row.cells[j+1], wd)
        write_cell(row.cells[4], "", valign="top"); set_col_width(row.cells[4], dim_cols[4])

    # Overall row
    ov_row = dim_table.add_row()
    write_cell(ov_row.cells[0], "OVERALL CASE RISK RATING", bold=True,
               color_hex=C.WHITE, bg_hex=C.NAVY); set_col_width(ov_row.cells[0], dim_cols[0])
    write_cell(ov_row.cells[1], "LOW",    bold=True, color_hex=C.GREEN_TX, bg_hex=C.GREEN_BG,
               align=WD_ALIGN_PARAGRAPH.CENTER); set_col_width(ov_row.cells[1], dim_cols[1])
    write_cell(ov_row.cells[2], "MEDIUM", bold=True, color_hex=C.AMBER_TX, bg_hex=C.AMBER_BG,
               align=WD_ALIGN_PARAGRAPH.CENTER); set_col_width(ov_row.cells[2], dim_cols[2])
    write_cell(ov_row.cells[3], "HIGH",   bold=True, color_hex=C.RED_TX,   bg_hex=C.RED_BG,
               align=WD_ALIGN_PARAGRAPH.CENTER); set_col_width(ov_row.cells[3], dim_cols[3])
    write_cell(ov_row.cells[4], "", valign="top"); set_col_width(ov_row.cells[4], dim_cols[4])

    doc.add_paragraph()
    narrative_box(doc, "Overall Risk Rationale (narrative justification for rating selected)", height_rows=4)
    doc.add_paragraph()

    label_input_with_data(doc, [
        ("Enhanced Due Diligence Required?",          "Yes / No — if Yes, specify measures", False),
        ("EDD Measures Undertaken / Planned",         "", False),
        ("SAR Required?",                             "Yes / No / Under Review — if Yes, note date submitted (do not record NCA reference here)", False),
        ("Ongoing Monitoring Frequency",              "e.g. At key milestones / Quarterly / 6-monthly / Annually", False),
        ("Next Scheduled Review Date",                "", False),
    ], label_cm=5.5)
    doc.add_paragraph()

    # ── SECTION 9 — IP SIGN-OFF & AMENDMENT LOG ──────────────────────────────
    section_banner(doc, 9, "IP Sign-Off and Amendment Log")
    guidance(doc, "The IP/Officeholder must sign off this assessment. Record all subsequent amendments.")

    so_cols = [CONTENT_CM / 3] * 3
    so_table = doc.add_table(rows=2, cols=3)
    so_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_row(so_table,
            [("IP / Officeholder Name", so_cols[0]),
             ("Signature", so_cols[1]),
             ("Date", so_cols[2])],
            bg=C.LIGHT_BLUE, text_color=C.BLACK)
    for i, wd in enumerate(so_cols):
        write_cell(so_table.cell(1, i), "", valign="top")
        set_col_width(so_table.cell(1, i), wd)

    doc.add_paragraph()
    p = doc.add_paragraph("Amendment Log")
    p.runs[0].bold = True; p.runs[0].font.name = "Arial"; p.runs[0].font.size = Pt(10)

    am_log_cols = [
        {"label": "Date",                "width": 2.0},
        {"label": "Amended By",          "width": 3.0},
        {"label": "Section(s) Amended",  "width": 5.5},
        {"label": "Reason for Amendment","width": 5.42},
    ]
    al_widths = [cd["width"] for cd in am_log_cols]
    al_table = doc.add_table(rows=0, cols=4)
    al_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_row(al_table, [(cd["label"], cd["width"]) for cd in am_log_cols],
            bg=C.LIGHT_BLUE, text_color=C.BLACK)
    for _ in range(5):
        blank_row(al_table, al_widths)

    doc.add_paragraph()
    guidance(doc,
             "This document is confidential and constitutes a record of AML/CFT checks undertaken "
             "in accordance with the Money Laundering, Terrorist Financing and Transfer of Funds "
             "(Information on the Payer) Regulations 2017 (as amended). Retain on the case file "
             "for a minimum of five years from the end of the business relationship.")

    return doc


# ─── Interactive prompts ──────────────────────────────────────────────────────

def prompt(label: str, default: str = "") -> str:
    if default:
        val = input(f"  {label} [{default}]: ").strip()
        return val if val else default
    else:
        return input(f"  {label}: ").strip()


def clean_company_number(raw: str) -> str:
    """Normalise company number: uppercase, zero-pad to 8 chars."""
    s = raw.strip().upper().replace(" ", "")
    # Scottish / NI prefixes: SC, NI, etc.
    prefix = re.match(r'^([A-Z]{2})', s)
    if prefix:
        digits = s[2:]
        return prefix.group(1) + digits.zfill(6)
    return s.zfill(8)


# ─── Main ─────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Generate a pre-filled AML Case Risk Assessment")
    parser.add_argument("--company", help="Companies House registration number")
    parser.add_argument("--api-key", help="Companies House API key (or set CH_API_KEY env var)")
    parser.add_argument("--output",  help="Output folder path")
    args = parser.parse_args()

    print("\n" + "="*60)
    print("  CASE-SPECIFIC RISK ASSESSMENT GENERATOR")
    print("  Giant Compliance Ltd  |  AML/CFT Tools")
    print("="*60 + "\n")

    # API key
    api_key = args.api_key or os.environ.get("CH_API_KEY", "")
    if not api_key:
        api_key = prompt("Companies House API key").strip()
    if not api_key:
        print("ERROR: API key required. Set CH_API_KEY environment variable or use --api-key.")
        sys.exit(1)

    # Company number
    raw_cn = args.company or prompt("Company Registration Number")
    company_number = clean_company_number(raw_cn)
    print(f"\n  Fetching data for company number: {company_number}")

    # Fetch from CH
    print("  [1/4] Company profile...")
    company_data = fetch_company(company_number, api_key)
    if not company_data.get("name"):
        print(f"\nERROR: Company {company_number} not found on Companies House.")
        sys.exit(1)
    print(f"        Found: {company_data['name']}")

    print("  [2/4] Officers (directors)...")
    officers = fetch_officers(company_number, api_key)
    print(f"        {len(officers['current'])} current, {len(officers['resigned'])} resigned officers")

    print("  [3/4] Persons with Significant Control...")
    pscs = fetch_pscs(company_number, api_key)
    print(f"        {len(pscs)} PSC(s) found")

    print("  [4/4] Filing history...")
    filing = fetch_filing_count(company_number, api_key)
    print(f"        {filing['total']} total filings, {filing['accounts']} accounts filings")

    # Case details from user
    print(f"\n  --- Case Details (press Enter to leave blank) ---")
    case_inputs = {
        "case_ref":    prompt("Case Reference (e.g. CVL001)"),
        "appt_type":   prompt("Appointment Type", "CVL"),
        "appt_date":   prompt("Date of Appointment (DD/MM/YYYY)"),
        "ip_name":     prompt("IP / Officeholder Name"),
        "ip_licence":  prompt("IP Licence Number"),
        "assessed_by": prompt("Assessed By"),
    }

    # Build document
    print("\n  Building document...")
    doc = build_document(company_data, officers, pscs, filing, case_inputs)

    # Output path
    safe_name = re.sub(r'[^\w\s-]', '', company_data["name"]).strip().replace(" ", "_")[:40]
    filename   = f"Case_RA_{company_number}_{safe_name}.docx"

    if args.output:
        out_dir = args.output
    else:
        # Default: same folder as this script
        out_dir = os.path.dirname(os.path.abspath(__file__))

    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, filename)
    doc.save(out_path)

    print("\n  Done! Document saved to:")
    print(f"  {out_path}\n")
    print("  [CH] labels = auto-populated from Companies House. Please verify.")


if __name__ == "__main__":
    main()
